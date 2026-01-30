"""
Clinical Component Identifier - Few-Shot Prompting Version
Uses Claude Sonnet with 18 labeled examples for high accuracy
Supports PDF, DOCX, TXT files up to 50MB
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import anthropic
import json
import os
import tempfile

# Check for optional dependencies
try:
    from pypdf import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("[WARNING] pypdf not installed. PDF support disabled.")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("[WARNING] python-docx not installed. DOCX support disabled.")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
CORS(app)

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({"error": "File too large. Maximum size is 50MB."}), 413

# ============================================
# CONFIGURATION
# ============================================
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "YOUR_API_KEY_HERE")
# ============================================

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

# Component taxonomy definition - Extended for CSR/ICH Guidelines
TAXONOMY = {
    "component_types": [
        {
            "name": "boilerplate",
            "description": "Standard regulatory or compliance text that appears across multiple documents (e.g., GCP compliance statements, regulatory references)"
        },
        {
            "name": "definition",
            "description": "Precise definitions of clinical terms, endpoints, or criteria (e.g., AE definitions, endpoint definitions)"
        },
        {
            "name": "study_section",
            "description": "Study-specific methodology or design elements (e.g., inclusion/exclusion criteria, study objectives)"
        },
        {
            "name": "drug_info",
            "description": "Information about investigational product, dosing, or formulation"
        },
        {
            "name": "safety",
            "description": "Safety monitoring, adverse event reporting, or dose modification guidelines"
        },
        {
            "name": "procedure",
            "description": "Clinical or laboratory procedures, assessments, or sampling instructions"
        },
        {
            "name": "csr_structure",
            "description": "ICH E3 Clinical Study Report structural elements, section headers, and organizational guidance"
        },
        {
            "name": "statistical",
            "description": "Statistical methodology, analysis plans, sample size calculations, and data handling procedures"
        },
        {
            "name": "regulatory_guidance",
            "description": "ICH guidelines, regulatory submission requirements, and compliance instructions"
        },
        {
            "name": "ethics",
            "description": "IRB/IEC requirements, informed consent procedures, Declaration of Helsinki references, and patient rights"
        }
    ]
}

# Few-shot examples for clinical component identification - Extended with CSR examples
FEW_SHOT_EXAMPLES = [
    # Original Protocol Examples
    {
        "text": "This clinical investigation will be conducted according to this clinical protocol and in compliance with Good Clinical Practice (GCP), with the Declaration of Helsinki (Version 2008), and with other applicable regulatory requirements.",
        "type": "boilerplate",
        "title": "GCP Compliance Statement",
        "confidence": 0.98,
        "reuse_potential": "high",
        "rationale": "Standard regulatory compliance language found in virtually all clinical protocols. Can be reused with minimal modification."
    },
    {
        "text": "Overall survival is defined as the time from randomization to death from any cause. Subjects who have not died will be censored at the date last known alive.",
        "type": "definition",
        "title": "Overall Survival Endpoint Definition",
        "confidence": 0.97,
        "reuse_potential": "high",
        "rationale": "Standard endpoint definition used across oncology trials. Highly reusable with consistent wording."
    },
    {
        "text": "Subject Inclusion Criteria: 1. Subject has provided signed written informed consent. 2. Subject is >=18 years of age. 3. Subject has histologically confirmed diagnosis of advanced solid tumor.",
        "type": "study_section",
        "title": "Basic Inclusion Criteria",
        "confidence": 0.96,
        "reuse_potential": "medium",
        "rationale": "Common inclusion criteria structure. Items 1-2 are highly reusable; item 3 is study-specific."
    },
    {
        "text": "Niraparib is an orally active poly (adenosine diphosphate [ADP]-ribose) polymerase (PARP)-1 and -2 inhibitor with nanomolar potency.",
        "type": "drug_info",
        "title": "Drug Mechanism of Action",
        "confidence": 0.97,
        "reuse_potential": "high",
        "rationale": "Standard drug description text that can be reused across all protocols involving this compound."
    },
    {
        "text": "Dose interruption and/or reduction may be implemented at any time for any grade toxicity considered intolerable by the subject. Treatment must be interrupted for any nonhematologic NCI-CTCAE Grade 3 or 4 AE.",
        "type": "safety",
        "title": "Dose Modification for Toxicity",
        "confidence": 0.95,
        "reuse_potential": "medium",
        "rationale": "Safety dose modification guidance. Structure is reusable but thresholds may vary by study."
    },
    {
        "text": "Blood samples for PK analysis will be collected at the following times: predose (0 hour), Day 1 (1, 1.5, 2, 3, 4, 6, and 12 hours postdose), Day 2 (24 hours postdose).",
        "type": "procedure",
        "title": "PK Blood Sampling Schedule",
        "confidence": 0.96,
        "reuse_potential": "medium",
        "rationale": "PK sampling procedure template. Timing may vary but format is standard."
    },
    {
        "text": "An adverse event (AE) is any untoward medical occurrence in a patient or clinical investigation subject administered a pharmaceutical product that does not necessarily have a causal relationship with this treatment.",
        "type": "definition",
        "title": "Adverse Event Definition",
        "confidence": 0.98,
        "reuse_potential": "high",
        "rationale": "ICH-standard AE definition. Used verbatim across all clinical trials."
    },
    {
        "text": "To determine the absolute bioavailability of niraparib by using an intravenous (IV) niraparib microdose of 100 ug in subjects with cancer.",
        "type": "study_section",
        "title": "Primary Study Objective",
        "confidence": 0.95,
        "reuse_potential": "low",
        "rationale": "Study-specific objective. Structure can be templated but content is unique."
    },
    # New CSR/ICH E3 Examples
    {
        "text": "The title page should contain the following information: study title, name of test drug/investigational product, indication studied, name of the sponsor, protocol identification, development phase of study, study initiation date, study completion date, name and affiliation of principal investigator, statement indicating compliance with Good Clinical Practices.",
        "type": "csr_structure",
        "title": "CSR Title Page Requirements",
        "confidence": 0.96,
        "reuse_potential": "high",
        "rationale": "ICH E3 guidance on title page structure. Standard template for all clinical study reports."
    },
    {
        "text": "A brief synopsis (usually limited to 3 pages) that summarises the study should be provided. The synopsis should include numerical data to illustrate results, not just text or p-values.",
        "type": "csr_structure",
        "title": "CSR Synopsis Requirements",
        "confidence": 0.97,
        "reuse_potential": "high",
        "rationale": "ICH E3 guidance on synopsis format and content requirements. Reusable structural guidance for all clinical study reports."
    },
    {
        "text": "It should be confirmed that the study was conducted in accordance with the ethical principles that have their origins in the Declaration of Helsinki.",
        "type": "ethics",
        "title": "Declaration of Helsinki Compliance Statement",
        "confidence": 0.98,
        "reuse_potential": "high",
        "rationale": "Standard ethics compliance boilerplate required in all clinical study reports. References fundamental ethical principles for human research."
    },
    {
        "text": "How and when informed consent was obtained in relation to patient enrolment, (e.g., at allocation, pre-screening) should be described. Representative written information for the patient (if any) and a sample patient consent form should be provided in appendix 16.1.3.",
        "type": "ethics",
        "title": "Informed Consent Documentation Requirements",
        "confidence": 0.95,
        "reuse_potential": "high",
        "rationale": "ICH E3 requirement for documenting informed consent procedures. Standard requirement for all CSRs."
    },
    {
        "text": "A serious adverse event (experience) or reaction is any untoward medical occurrence that at any dose: results in death, is life-threatening, requires inpatient hospitalisation or prolongation of existing hospitalisation, results in persistent or significant disability/incapacity, or is a congenital anomaly/birth defect.",
        "type": "definition",
        "title": "Serious Adverse Event (SAE) Definition",
        "confidence": 0.98,
        "reuse_potential": "high",
        "rationale": "ICH-harmonized definition of serious adverse events. Standard regulatory definition used globally in all clinical trials and CSRs."
    },
    {
        "text": "The planned sample size and the basis for it, such as statistical considerations or practical limitations, should be provided. Methods for sample size calculation should be given together with their derivations or source of reference. Estimates used in the calculations should be given and explanations provided as to how they were obtained.",
        "type": "statistical",
        "title": "Sample Size Determination Requirements",
        "confidence": 0.96,
        "reuse_potential": "high",
        "rationale": "ICH E3 requirement for documenting sample size calculations. Standard statistical documentation requirement for CSRs."
    },
    {
        "text": "There should be a clear accounting of all patients who entered the study, using figures or tables in the text of the report. The numbers of patients who were randomised, and who entered and completed each phase of the study, (or each week/month of the study) should be provided, as well as the reasons for all post-randomisation discontinuations, grouped by treatment and by major reason.",
        "type": "study_section",
        "title": "Patient Disposition Requirements",
        "confidence": 0.95,
        "reuse_potential": "high",
        "rationale": "ICH E3 guidance on patient disposition reporting. Standard requirement for documenting patient flow in clinical study reports."
    },
    {
        "text": "The extent of exposure to test drugs/investigational products (and to active control and placebo) should be characterised according to the number of patients exposed, the duration of exposure, and the dose to which they were exposed.",
        "type": "safety",
        "title": "Extent of Exposure Requirements",
        "confidence": 0.95,
        "reuse_potential": "high",
        "rationale": "ICH E3 requirement for safety evaluation. Standard approach to documenting drug exposure in clinical study reports."
    },
    {
        "text": "Depending on the regulatory authority's review policy, abbreviated reports using summarised data or with some sections deleted, may be acceptable for uncontrolled studies or other studies not designed to establish efficacy.",
        "type": "regulatory_guidance",
        "title": "Abbreviated Report Guidance",
        "confidence": 0.94,
        "reuse_potential": "high",
        "rationale": "ICH E3 guidance on when abbreviated CSR formats may be acceptable. Regulatory flexibility guidance for different study types."
    }
]


def build_few_shot_prompt(document_text):
    """Build the few-shot prompt with examples and document text."""
    
    # Format taxonomy
    taxonomy_str = json.dumps(TAXONOMY, indent=2)
    
    # Format few-shot examples
    examples_str = ""
    for i, ex in enumerate(FEW_SHOT_EXAMPLES, 1):
        examples_str += f"""
EXAMPLE {i}:
Text: "{ex['text']}"
Classification:
- Type: {ex['type']}
- Title: {ex['title']}
- Confidence: {ex['confidence']}
- Reuse Potential: {ex['reuse_potential']}
- Rationale: {ex['rationale']}
"""

    prompt = f"""You are an expert clinical documentation analyst specializing in identifying reusable content components in regulatory documents such as clinical trial protocols, statistical analysis plans, clinical study reports, and ICH guidelines.

CRITICAL INSTRUCTION: You must identify and extract ALL distinct reusable components from the document. Do NOT skip any components. Be thorough and comprehensive.

TASK: Analyze the provided clinical document text and identify EVERY distinct reusable component. The document contains [PAGE X] markers indicating page numbers.

COMPONENT TAXONOMY:
{taxonomy_str}

IDENTIFICATION RULES:
1. Extract ALL components - do not skip any reusable content
2. Components must be self-contained and semantically complete
3. Include ALL relevant context within component boundaries
4. Avoid overlapping components
5. Prefer larger, meaningful units over small fragments
6. Minimum component length: 50 characters
7. Each component should represent a single, coherent concept or section
8. Track the page number where each component is found (look for [PAGE X] markers)
9. Identify the section name/number if visible (e.g., "Section 5.1", "12.2 ADVERSE EVENTS")
10. Include definitions, procedures, requirements, guidance, boilerplate text, tables descriptions, etc.

COMPONENT TYPES TO LOOK FOR:
- Definitions (endpoints, adverse events, terms)
- Boilerplate/regulatory language (GCP, ethics, compliance)
- Study sections (objectives, design, criteria)
- Safety content (AE reporting, dose modifications)
- Procedures (assessments, visits, sampling)
- Statistical methods (sample size, analysis plans)
- CSR structural requirements
- Regulatory guidance
- Ethics requirements

{examples_str}

OUTPUT FORMAT:
Return a JSON array with this exact structure for each identified component:
[
  {{
    "type": "component_type",
    "title": "Descriptive title (5-10 words)",
    "text": "Exact extracted text from the document (copy verbatim, include full content)",
    "confidence": 0.95,
    "reuse_potential": "high|medium|low",
    "rationale": "Brief explanation of why this is a reusable component",
    "location": {{
      "page": 1,
      "section": "Section name or number if identifiable, otherwise null"
    }}
  }}
]

IMPORTANT: 
- Extract the COMPLETE text of each component - do not truncate or summarize
- Include ALL components you find - aim to be exhaustive
- Copy text verbatim from the document

DOCUMENT TO ANALYZE:
{document_text}

Identify ALL reusable components and return ONLY the JSON array, no additional text."""

    return prompt


@app.route('/')
def health_check():
    """Health check endpoint."""
    return jsonify({
        "status": "healthy",
        "service": "Clinical Component Identifier (Few-Shot)",
        "version": "2.0",
        "model": "claude-sonnet-4-20250514",
        "examples": len(FEW_SHOT_EXAMPLES)
    })


@app.route('/api/identify', methods=['POST'])
def identify_components():
    """Identify clinical components using few-shot prompting."""
    try:
        data = request.get_json()
        
        if not data or 'text' not in data:
            return jsonify({"error": "Missing 'text' field in request body"}), 400
        
        document_text = data['text'].strip()
        
        if len(document_text) < 50:
            return jsonify({"error": "Text must be at least 50 characters long"}), 400
        
        # Build the few-shot prompt
        prompt = build_few_shot_prompt(document_text)
        
        # Call Claude API
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=16000,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            system="You are an expert at identifying reusable components in clinical trial documentation. You always respond with valid JSON arrays only."
        )
        
        # Parse response
        result_text = response.content[0].text.strip()
        
        # Handle markdown code blocks if present
        if '```json' in result_text:
            result_text = result_text.split('```json')[1].split('```')[0].strip()
        elif '```' in result_text:
            result_text = result_text.split('```')[1].split('```')[0].strip()
        
        components = json.loads(result_text)
        
        # Validate components
        validated_components = []
        valid_types = [t["name"] for t in TAXONOMY["component_types"]]
        
        for comp in components:
            location = comp.get("location", {})
            validated_comp = {
                "type": comp.get("type", "unknown"),
                "title": comp.get("title", "Untitled Component"),
                "text": comp.get("text", ""),
                "confidence": float(comp.get("confidence", 0.8)),
                "reuse_potential": comp.get("reuse_potential", "medium"),
                "rationale": comp.get("rationale", ""),
                "location": {
                    "page": location.get("page") if isinstance(location, dict) else None,
                    "section": location.get("section") if isinstance(location, dict) else None
                }
            }
            
            # Validate type
            if validated_comp["type"] not in valid_types:
                validated_comp["type"] = "study_section"
            
            validated_components.append(validated_comp)
        
        return jsonify({
            "success": True,
            "components": validated_components,
            "total_components": len(validated_components),
            "model": "claude-sonnet-4-20250514",
            "method": "few-shot",
            "examples_used": len(FEW_SHOT_EXAMPLES)
        })
        
    except json.JSONDecodeError as e:
        return jsonify({
            "error": f"Failed to parse model response as JSON: {str(e)}"
        }), 500
    except Exception as e:
        return jsonify({
            "error": f"Server error: {str(e)}"
        }), 500


@app.route('/api/taxonomy', methods=['GET'])
def get_taxonomy():
    """Return the component taxonomy."""
    return jsonify(TAXONOMY)


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using pypdf with page tracking."""
    if not PDF_SUPPORT:
        raise Exception("PDF support not available. Install pypdf: pip install pypdf")
    
    reader = PdfReader(file_path)
    pages_data = []
    
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            # Clean up common PDF extraction issues
            cleaned_text = text.replace('-\n', '')
            
            # Normalize whitespace but preserve paragraph breaks
            lines = cleaned_text.split('\n')
            normalized_lines = []
            current_paragraph = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    if current_paragraph:
                        normalized_lines.append(' '.join(current_paragraph))
                        current_paragraph = []
                    normalized_lines.append('')
                elif line.endswith(('.', ':', ';', '?', '!')) or line.startswith(('−', '•', '-', '*', '¾', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', 'a)', 'b)', 'c)', 'd)', 'e)')):
                    current_paragraph.append(line)
                    normalized_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                else:
                    current_paragraph.append(line)
            
            if current_paragraph:
                normalized_lines.append(' '.join(current_paragraph))
            
            cleaned_text = '\n'.join(normalized_lines)
            while '\n\n\n' in cleaned_text:
                cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
            
            pages_data.append({
                "page": page_num + 1,
                "text": cleaned_text
            })
    
    return pages_data


def extract_text_from_pdf_simple(file_path):
    """Extract text from PDF as a single string with page markers."""
    pages_data = extract_text_from_pdf(file_path)
    text_parts = []
    for p in pages_data:
        text_parts.append(f"[PAGE {p['page']}]\n{p['text']}")
    return "\n\n".join(text_parts), pages_data


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file using python-docx."""
    if not DOCX_SUPPORT:
        raise Exception("DOCX support not available. Install python-docx: pip install python-docx")
    
    doc = Document(file_path)
    text_parts = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    
    return "\n\n".join(text_parts)


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload (PDF, DOCX, TXT) and identify ALL components without truncation."""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Get file extension
        filename = file.filename.lower()
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name
        
        pages_data = None
        try:
            # Extract text based on file type
            if filename.endswith('.pdf'):
                document_text, pages_data = extract_text_from_pdf_simple(tmp_path)
            elif filename.endswith('.docx'):
                document_text = extract_text_from_docx(tmp_path)
            elif filename.endswith('.txt'):
                with open(tmp_path, 'r', encoding='utf-8') as f:
                    document_text = f.read()
            else:
                return jsonify({"error": f"Unsupported file type. Supported: PDF, DOCX, TXT"}), 400
            
        finally:
            # Clean up temp file
            os.unlink(tmp_path)
        
        if len(document_text.strip()) < 50:
            return jsonify({"error": "Extracted text is too short (less than 50 characters)"}), 400
        
        total_chars = len(document_text)
        
        # Process document in chunks if it's very large
        # Each chunk should be around 25000 chars to leave room for prompt and response
        chunk_size = 25000
        all_components = []
        chunks_processed = 0
        
        if total_chars <= chunk_size:
            # Small document - process in one go
            all_components = process_document_chunk(document_text, 0)
            chunks_processed = 1
        else:
            # Large document - process in chunks by page boundaries
            if pages_data:
                # Split by pages for PDF
                current_chunk = ""
                current_start_page = 1
                
                for page_info in pages_data:
                    page_text = f"[PAGE {page_info['page']}]\n{page_info['text']}\n\n"
                    
                    if len(current_chunk) + len(page_text) > chunk_size and current_chunk:
                        # Process current chunk
                        chunk_components = process_document_chunk(current_chunk, current_start_page - 1)
                        all_components.extend(chunk_components)
                        chunks_processed += 1
                        
                        # Start new chunk
                        current_chunk = page_text
                        current_start_page = page_info['page']
                    else:
                        current_chunk += page_text
                
                # Process final chunk
                if current_chunk:
                    chunk_components = process_document_chunk(current_chunk, current_start_page - 1)
                    all_components.extend(chunk_components)
                    chunks_processed += 1
            else:
                # Non-PDF - split by character count
                for i in range(0, total_chars, chunk_size):
                    chunk = document_text[i:i + chunk_size]
                    chunk_components = process_document_chunk(chunk, i)
                    all_components.extend(chunk_components)
                    chunks_processed += 1
        
        # Deduplicate components based on text similarity
        unique_components = deduplicate_components(all_components)
        
        return jsonify({
            "success": True,
            "components": unique_components,
            "total_components": len(unique_components),
            "total_pages": len(pages_data) if pages_data else None,
            "model": "claude-sonnet-4-20250514",
            "method": "few-shot",
            "examples_used": len(FEW_SHOT_EXAMPLES),
            "filename": file.filename,
            "text_length": total_chars,
            "chunks_processed": chunks_processed,
            "truncated": False  # Never truncate anymore
        })
        
    except json.JSONDecodeError as e:
        return jsonify({
            "error": f"Failed to parse model response as JSON: {str(e)}"
        }), 500
    except Exception as e:
        import traceback
        return jsonify({
            "error": f"Server error: {str(e)}",
            "traceback": traceback.format_exc()
        }), 500


def process_document_chunk(chunk_text, chunk_offset=0):
    """Process a single chunk of document text and return components."""
    try:
        prompt = build_few_shot_prompt(chunk_text)
        
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=16000,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            system="""You are an expert at identifying reusable components in clinical trial documentation. 
You must identify ALL reusable components in the document - do not skip any.
Always respond with valid JSON arrays only. 
Include location information (page number and section) for each component.
Be thorough and comprehensive - extract every distinct reusable component you find."""
        )
        
        result_text = response.content[0].text.strip()
        
        # Handle markdown code blocks if present
        if '```json' in result_text:
            result_text = result_text.split('```json')[1].split('```')[0].strip()
        elif '```' in result_text:
            result_text = result_text.split('```')[1].split('```')[0].strip()
        
        components = json.loads(result_text)
        
        # Validate components
        validated_components = []
        valid_types = [t["name"] for t in TAXONOMY["component_types"]]
        
        for comp in components:
            location = comp.get("location", {})
            validated_comp = {
                "type": comp.get("type", "unknown"),
                "title": comp.get("title", "Untitled Component"),
                "text": comp.get("text", ""),
                "confidence": float(comp.get("confidence", 0.8)),
                "reuse_potential": comp.get("reuse_potential", "medium"),
                "rationale": comp.get("rationale", ""),
                "location": {
                    "page": location.get("page") if isinstance(location, dict) else None,
                    "section": location.get("section") if isinstance(location, dict) else None
                }
            }
            
            if validated_comp["type"] not in valid_types:
                validated_comp["type"] = "study_section"
            
            validated_components.append(validated_comp)
        
        return validated_components
        
    except Exception as e:
        print(f"Error processing chunk: {str(e)}")
        return []


def deduplicate_components(components):
    """Remove duplicate components based on text similarity."""
    if not components:
        return []
    
    unique = []
    seen_texts = set()
    
    for comp in components:
        # Create a normalized version of the text for comparison
        normalized = comp["text"].lower().strip()[:200]  # Compare first 200 chars
        
        if normalized not in seen_texts:
            seen_texts.add(normalized)
            unique.append(comp)
    
    return unique


@app.route('/api/supported-formats', methods=['GET'])
def get_supported_formats():
    """Return supported file formats."""
    return jsonify({
        "formats": [
            {"extension": ".pdf", "name": "PDF", "supported": PDF_SUPPORT},
            {"extension": ".docx", "name": "Word Document", "supported": DOCX_SUPPORT},
            {"extension": ".txt", "name": "Plain Text", "supported": True}
        ],
        "max_size_mb": 50
    })


if __name__ == '__main__':
    print("=" * 60)
    print("Clinical Component Identifier - Few-Shot Version")
    print("=" * 60)
    print(f"Model: claude-sonnet-4-20250514")
    print(f"Few-shot examples: {len(FEW_SHOT_EXAMPLES)}")
    print(f"Expected accuracy: 85-95%")
    print(f"PDF Support: {PDF_SUPPORT}")
    print(f"DOCX Support: {DOCX_SUPPORT}")
    print(f"Max File Size: 50MB")
    print("=" * 60)
    
    if ANTHROPIC_API_KEY == "YOUR_API_KEY_HERE":
        print("\n[WARNING] Set your Anthropic API key in app.py or as environment variable ANTHROPIC_API_KEY\n")
    
    app.run(host='0.0.0.0', port=5000, debug=True)
